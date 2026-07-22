import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Clearcase Deck-Transparent Dispute & Chargeback Resolution Engine, Powered by Program Synthesis', 0)
    
    # Subtitle
    doc.add_paragraph('CodeStreet Idea Submission — Track: Frictionless Dispute & Chargeback Resolution')
    doc.add_paragraph('Team: Pranay | College: Chitkara University')
    
    doc.add_heading('The Problem', level=1)
    doc.add_paragraph("Right now, resolving a disputed charge takes weeks. Evidence gets requested, a merchant responds (or doesn't), someone eventually makes a call, and neither the card member nor the merchant really understands why the decision landed the way it did. That's slow, it's opaque, and it erodes trust on both sides regardless of who \"wins\" the dispute. Our solution is built around fixing both problems at once speed and transparency  rather than trading one for the other.")
    
    doc.add_heading('1. Why This Effectively Solves the Problem', level=1)
    doc.add_paragraph("We've built a system called Brain2  a neuro-symbolic program synthesis engine with three working parts:")
    
    doc.add_paragraph("Eyes — a small language model that reads raw input (a card member's claim, a merchant's response, shipment records, chargeback codes) and converts it into clean, structured data.", style='List Bullet')
    doc.add_paragraph("Brain — the synthesis core. Instead of training a black-box classifier, it searches for a small, explicit rule that correctly explains a set of past resolved disputes. The output is logic, not a probability score.", style='List Bullet')
    doc.add_paragraph("Mouth — a small language model that turns that synthesized logic back into a plain-English explanation for the card member and merchant.", style='List Bullet')
    
    doc.add_paragraph("Here's how each of the five stated tasks gets addressed:")
    
    doc.add_paragraph("1. Auto-collect and parse transaction evidence. Eyes handles this directly — reading card member claims, merchant responses, shipment records, and chargeback codes, and converting all of it into one structured format.")
    doc.add_paragraph("2. Fair-weighing model across card member and merchant. Brain synthesizes the decision rule, using the corroboration-weighted schema described in the innovation section below — evidence is weighed by how well it's corroborated, never by which party supplied it.")
    doc.add_paragraph("3. Interface for submission and status tracking. Both card members and merchants get a shared case view: a submission form for evidence (documents, photos, tracking numbers, written statements), and a live status tracker with clear stages — evidence received → under review → resolved — so neither side is left wondering where things stand. When a case resolves, both parties see the same plain-English explanation from Mouth, not two different summaries.")
    doc.add_paragraph("4. Transparent reasoning layer. Mouth translates Brain's synthesized rule into that plain-English explanation, shown to both parties at resolution.")
    doc.add_paragraph("5. Test and optimize the pipeline. We'd validate on three axes, not just one: accuracy (does the synthesized rule's decision match real historical outcomes on a held-out set of resolved disputes), speed (time from evidence submission to resolution, end to end), and fairness (a counterfactual audit — swap which party \"said\" a given piece of evidence while holding the evidence itself constant, and confirm the decision doesn't change). Any case that gets successfully appealed or overturned on human review feeds back into the training set, so the rule improves over time rather than staying static.")
    
    doc.add_paragraph("We're not bending an unrelated project to fit this problem — the underlying architecture already does the heavy lifting for tasks 1, 2, and 4. Tasks 3 and 5 are scoped, buildable engineering work rather than architectural risk, which is exactly the kind of thing a small team can execute confidently in a hackathon timeframe.")
    
    doc.add_page_break()
    
    doc.add_heading('2. The Innovation Behind This Idea', level=1)
    doc.add_paragraph("The obvious way to solve this problem is a classifier: feed it dispute and evidence data, get back an approve/deny score. It would probably work fine — and it's exactly what most other teams will likely pitch. It also inherits the same core problem it's replacing: a decision nobody can actually explain, just automated instead of manual.")
    doc.add_paragraph("Our approach is different in kind, not just in framing. Synthesis produces a rule, not a score — the explanation isn't a second model guessing at why the first one decided something, it's a direct readout of the actual logic that made the call. There's no approximation gap between what the system did and what it says it did.")
    doc.add_paragraph("The harder, more interesting problem is fairness between two conflicting parties, and here's our answer: Eyes parses both the card member's and merchant's evidence into the same structured schema, and every field is tagged by how corroborated it is — not by who provided it. A shipment tracking number that independently confirms delivery carries more weight than an unverified claim of \"it never arrived,\" regardless of which party said it. Brain never sees \"card member says X, merchant says Y\" — it sees one evidence table, weighted purely by evidential strength. Fairness becomes a property of the schema itself instead of a hand-tuned formula, which means there's no dial anyone can bias, intentionally or not.")
    
    doc.add_heading('3. Technical Feasibility', level=1)
    doc.add_paragraph("The core architecture — eyes, brain, and mouth, already exists and has been built and tested outside this hackathon, in a different domain. What's left for this challenge is genuinely scoped, not speculative: a corroboration-tagged evidence schema for disputes specifically, a dispute-relevant DSL vocabulary (chargeback codes, evidence categories, resolution types), a training set of resolved dispute examples, and the card member/merchant-facing interface.")
    
    doc.add_paragraph("Proposed stack:")
    doc.add_paragraph("Eyes: small LLM (Hugging Face), fine-tuned or prompted for structured extraction, similar to the spaCy/HF approach the brief itself suggests.", style='List Bullet')
    doc.add_paragraph("Brain: our existing synthesis framework, adapted with a dispute-specific DSL.", style='List Bullet')
    doc.add_paragraph("Mouth: small LLM, prompt-conditioned on the synthesized rule.", style='List Bullet')
    doc.add_paragraph("Interface: React frontend, FastAPI backend, PostgreSQL for case and audit storage.", style='List Bullet')
    doc.add_paragraph("Deployment: containerized, cloud-hosted (AWS).", style='List Bullet')
    
    doc.add_paragraph("None of this requires exotic infrastructure, which is deliberate, a system meant to run on real dispute volume should be cheap and boring to operate, not a research prototype held together with duct tape.")
    
    doc.add_heading('4. Ease of Implementation', level=1)
    doc.add_paragraph("We're not starting from a blank page. The hardest part of a project like this designing and validating a working synthesis architecture is already done. What remains is domain adaptation: building the dispute-specific schema and DSL, assembling training examples, and wiring up a standard web interface. That's real work, but it's the kind of work a small team can execute confidently within a hackathon timeline, because the architectural risk is already retired.")
    doc.add_paragraph("We'd rather be upfront about this split, what exists versus what we're building for CodeStreet specifically than imply a finished product. Judges can tell the difference, and it's a stronger position anyway: it shows the team already climbed the hard technical curve before the clock started.")
    
    doc.add_heading('5. Scalability', level=1)
    doc.add_paragraph("The decision-making core is a small, synthesized rule not a large model running inference on every dispute. That means the system scales close to linearly with dispute volume without a proportional rise in compute cost. The LLM components (eyes and mouth) only handle the edges, reading input in and writing explanation out, which is the one place natural language actually needs to be understood or generated. A containerized deployment on standard cloud infrastructure means scaling to a higher dispute volume is an infrastructure decision, not a re-architecture.")
    
    doc.add_heading('6. Business Relevance', level=1)
    doc.add_paragraph("Faster, more transparent dispute resolution is directly relevant to two of AmEx's core relationships at once: card member trust and merchant trust. Card members currently wait weeks for a resolution they can't see the reasoning behind; merchants face the same opacity from the other side. A system that's both faster and legibly fair reduces friction for both audiences simultaneously, which is a rare combination, most efficiency gains in this space come at the cost of transparency, or vice versa.")
    
    doc.add_heading('7. Potential Impact', level=1)
    doc.add_paragraph("Directionally, the impact shows up in a few concrete places: dispute cycle time moving from weeks toward minutes for the resolvable majority of cases, fewer re-disputes and escalations because the reasoning is visible rather than asserted, and lower operational load on human reviewers, who get freed up for genuinely ambiguous cases instead of routine ones. We haven't fabricated specific percentage figures here because we don't have access to AmEx's actual dispute volume or current resolution-time data, if selected, validating these numbers against real (or provided sample) data would be one of our first steps, and we'd rather say that than invent statistics that sound impressive but aren't grounded in anything.")
    
    doc.save('CodeStreet_ClearCase_Submission_Exact.docx')
    print("Word document created successfully with exact formatting.")

if __name__ == '__main__':
    main()
